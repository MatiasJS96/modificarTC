import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document
from datetime import timedelta
import re
import os

def tc_to_timedelta(tc, fps):
    try:
        h, m, s, f = map(int, tc.strip().split(':'))
        total_seconds = h * 3600 + m * 60 + s + f / fps
        return timedelta(seconds=total_seconds)
    except Exception:
        raise ValueError("Formato de TC inválido. Usá hh:mm:ss:ff")

def timedelta_to_tc(td, fps):
    total_seconds = td.total_seconds()
    h = int(total_seconds // 3600)
    m = int((total_seconds % 3600) // 60)
    s = int(total_seconds % 60)
    f = int(round((total_seconds - int(total_seconds)) * fps))
    if f >= fps:
        f = 0
        s += 1
        if s >= 60:
            s = 0
            m += 1
            if m >= 60:
                m = 0
                h += 1
    return f"{h:02}:{m:02}:{s:02}:{f:02}"

def ajustar_tc(doc, delta, fps):
    tc_pattern = re.compile(r'\b\d{2}:\d{2}:\d{2}:\d{2}\b')
    
    for para in doc.paragraphs:
        matches = tc_pattern.findall(para.text)
        new_text = para.text
        for match in matches:
            updated_tc = timedelta_to_tc(tc_to_timedelta(match, fps) + delta, fps)
            new_text = new_text.replace(match, updated_tc)
        para.text = new_text

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    matches = tc_pattern.findall(para.text)
                    new_text = para.text
                    for match in matches:
                        updated_tc = timedelta_to_tc(tc_to_timedelta(match, fps) + delta, fps)
                        new_text = new_text.replace(match, updated_tc)
                    para.text = new_text

def seleccionar_archivo():
    filepath = filedialog.askopenfilename(filetypes=[("Documentos Word", "*.docx")])
    if filepath:
        archivo_var.set(filepath)

def procesar():
    ruta = archivo_var.get()
    tc_original = tc_original_var.get()
    tc_nuevo = tc_nuevo_var.get()

    try:
        fps = float(fps_var.get())
        td_original = tc_to_timedelta(tc_original, fps)
        td_nuevo = tc_to_timedelta(tc_nuevo, fps)
        delta = td_nuevo - td_original
    except ValueError as e:
        messagebox.showerror("Error en TC", str(e))
        return

    if not ruta or not os.path.exists(ruta):
        messagebox.showerror("Error", "Seleccioná un archivo válido.")
        return

    try:
        doc = Document(ruta)
        ajustar_tc(doc, delta, fps)
        base = os.path.splitext(ruta)[0]
        salida = f"{base}_AJUSTADO_desde_{tc_original.replace(':', '-')}_a_{tc_nuevo.replace(':', '-')}_{fps}fps.docx"
        doc.save(salida)
        messagebox.showinfo("Éxito", f"Archivo guardado como:\n{salida}")
    except Exception as e:
        messagebox.showerror("Error", str(e))

# Interfaz
root = tk.Tk()
root.title("Ajustador de Timecodes por TC de Referencia")

archivo_var = tk.StringVar()
tc_original_var = tk.StringVar()
tc_nuevo_var = tk.StringVar()
fps_var = tk.StringVar(value="23.976")

tk.Label(root, text="Archivo DOCX:").pack(pady=(10,0))
tk.Entry(root, textvariable=archivo_var, width=60).pack(padx=10)
tk.Button(root, text="Seleccionar archivo", command=seleccionar_archivo).pack(pady=5)

tk.Label(root, text="TC original de referencia (ej. 01:00:00:00):").pack()
tk.Entry(root, textvariable=tc_original_var).pack()

tk.Label(root, text="Nuevo TC deseado (ej. 01:02:30:10):").pack()
tk.Entry(root, textvariable=tc_nuevo_var).pack()

tk.Label(root, text="Framerate (fps):").pack()
tk.OptionMenu(root, fps_var, "23.976", "24", "25", "29.97").pack()

tk.Button(root, text="Procesar", command=procesar).pack(pady=10)

root.mainloop()